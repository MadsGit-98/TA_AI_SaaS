"""
Resume Parsing Service

Per Constitution §4: Decoupled services located in project root services/ directory.

This service handles:
1. PDF text extraction using PyPDF2
2. Docx text extraction using python-docx
3. Confidential information filtering (PII redaction)
"""

import re
import hashlib
from pypdf import PdfReader
from docx import Document
from io import BytesIO
import phonenumbers


class ResumeParserService:
    """
    Service for extracting text from PDF and Docx resume files.
    """
    
    @staticmethod
    def extract_text_from_pdf(file_content: bytes) -> str:
        """
        Extract text from a PDF file.
        
        Args:
            file_content: Raw bytes of the PDF file
            
        Returns:
            Extracted text content
        """
        reader = PdfReader(BytesIO(file_content))
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text.strip()
    
    @staticmethod
    def extract_text_from_docx(file_content: bytes) -> str:
        """
        Extract text from a Docx file.

        Args:
            file_content: Raw bytes of the Docx file

        Returns:
            Extracted text content
        """
        doc = Document(BytesIO(file_content))
        text = ""
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        # Extract text from tables
        # Use a set to track exact cell values seen (not substring matching)
        seen_cells = set()
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text and cell_text not in seen_cells:
                        text += cell_text + "\n"
                        seen_cells.add(cell_text)
        return text.strip()
    
    @staticmethod
    def calculate_file_hash(file_content: bytes) -> str:
        """
        Calculate SHA-256 hash of file content for duplication detection.
        
        Args:
            file_content: Raw bytes of the file
            
        Returns:
            Hexadecimal hash string (64 characters)
        """
        return hashlib.sha256(file_content).hexdigest()


class ConfidentialInfoFilter:
    """
    Filter to redact confidential personal information from parsed resume text.
    
    Per specification FR-016: Confidential info (phone, email, addresses) must not
    be stored in parsed text for AI analysis, but contact info is stored separately
    for communication purposes.
    """
    
    # Email pattern
    EMAIL_PATTERN = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    
    # Phone pattern - comprehensive regex to capture various formats with extensions
    # This is used as a first pass, then validated with phonenumbers library
    PHONE_PATTERN = (
        r'(?:'
        # International format: +1 123 456 7890, +1-123-456-7890, +44 20 7946 0958
        r'\+\d{1,3}[-.\s]?\(?\d{1,4}\)?[-.\s]?\d{1,4}[-.\s]?\d{1,9}'
        r'|'
        # US format with country code: 1-123-456-7890, 1 (123) 456-7890
        r'1[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
        r'|'
        # US local format: (123) 456-7890, 123-456-7890, 123.456.7890, 123 456 7890
        r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
        r'|'
        # Alphanumeric toll-free: 1-800-FLOWERS, 1-888-CALL-NOW
        r'1-\d{3}-[A-Z]{4,10}'
        r')'
        # Optional extension: x123, ext. 456, extension 789
        r'(?:\s*(?:x|ext\.?|extension)\s*\d+)?'
    )
    
    # SSN pattern (matches XXX-XX-XXXX, XXX XX XXXX, or XXXXXXXXXX)
    SSN_PATTERN = r'\b\d{3}[-\s]?\d{2}[-\s]?\d{4}\b'

    # Month names for date patterns
    MONTH_NAMES = r'(?:January|February|March|April|May|June|July|August|September|October|November|December)'

    # Contextual prefix for DOB detection
    # Matches: DOB, Date of Birth, Born, Birthday, D.O.B. (with optional punctuation/spaces)
    DOB_CONTEXT_PREFIX = r'(?:DOB|Date\s+of\s+Birth|Born|Birthday|D\.?\s*O\.?\s*B\.?)[:.]?\s*'

    # Date of birth pattern - matches multiple formats ONLY when preceded by DOB context:
    # - DOB: MM/DD/YYYY or DOB: MM-DD-YYYY (US format)
    # - DOB: YYYY-MM-DD (ISO format)
    # - DOB: Month DD, YYYY or DOB: DD Month YYYY (written month format)
    DOB_PATTERN = (
        r'\b' + DOB_CONTEXT_PREFIX + r'(?:'
        # MM/DD/YYYY or MM-DD-YYYY
        r'(0[1-9]|1[0-2])[-/](0[1-9]|[12]\d|3[01])[-/](19|20)\d{2}'
        r'|'
        # YYYY-MM-DD (ISO)
        r'(19|20)\d{2}-(0[1-9]|1[0-2])-(0[1-9]|[12]\d|3[01])'
        r'|'
        # Month DD, YYYY (e.g., January 15, 1990)
        r'' + MONTH_NAMES + r'\s+(0[1-9]|[12]\d|3[01]),?\s+(19|20)\d{2}'
        r'|'
        # DD Month YYYY (e.g., 15 January 1990)
        r'(0[1-9]|[12]\d|3[01])\s+' + MONTH_NAMES + r',?\s+(19|20)\d{2}'
        r')\b'
    )
    
    # Street address patterns (enhanced for comprehensive coverage)
    # Comprehensive street suffix list
    STREET_SUFFIXES = (
        r'(?:Street|St|Avenue|Ave|Road|Rd|Boulevard|Blvd|Drive|Dr|Lane|Ln|'
        r'Court|Ct|Way|Place|Pl|Parkway|Pkwy|Circle|Cir|Terrace|Ter|Trail|Trl|'
        r'Highway|Hwy|Square|Sq|Plaza|Plz|Loop|Alley|Alcove|Pass|Path)'
    )

    # US State abbreviations (including DC)
    US_STATES = (
        r'(?:AL|AK|AZ|AR|CA|CO|CT|DE|FL|GA|HI|ID|IL|IN|IA|KS|KY|LA|ME|MD|'
        r'MA|MI|MN|MS|MO|MT|NE|NV|NH|NJ|NM|NY|NC|ND|OH|OK|OR|PA|RI|SC|SD|'
        r'TN|TX|UT|VT|VA|WA|WV|WI|WY|DC)'
    )

    # Canadian province abbreviations
    CA_PROVINCES = r'(?:AB|BC|MB|NB|NL|NS|NT|NU|ON|PE|QC|SK|YT)'

    # US Street Address pattern (street number + name + suffix)
    US_STREET_PATTERN = r'\b\d+\s+[A-Za-z\s]+?' + STREET_SUFFIXES + r'\b'

    # US City, State, ZIP pattern (with optional comma separators)
    US_CITY_ZIP_PATTERN = (
        r',?\s*[A-Za-z\s]+?,?\s*' + US_STATES + r'\s+\d{5}(?:-\d{4})?\b'
    )

    # Canadian Postal Code pattern (format: A1A 1A1 or A1A1A1)
    CANADIAN_POSTAL_CODE_PATTERN = r'[A-Z]\d[A-Z]\s?\d[A-Z]\d\b'

    # Canadian City, Province, Postal Code pattern
    CA_CITY_POSTAL_PATTERN = (
        r',?\s*[A-Za-z\s]+?,?\s*' + CA_PROVINCES + r'\s+' + 
        CANADIAN_POSTAL_CODE_PATTERN
    )

    # Apartment/Suite/Unit pattern
    APT_SUITE_UNIT_PATTERN = (
        r'(?:Apt\.?|Apartment|Suite|Unit|Lot|Space|#)\s*[A-Za-z0-9]+'
    )

    # PO Box pattern
    PO_BOX_PATTERN = r'\b(?:PO\s*Box|P\.?\s*O\.?\s*Box)\s*\d+\b'

    # Rural Route pattern
    RURAL_ROUTE_PATTERN = r'\b(?:RR|Rural\s*Route)\s*\d+\s*(?:Box|B)\s*\d+\b'

    # Full US Address pattern (combines street + city/state/zip + optional apt)
    US_ADDRESS_PATTERN = (
        US_STREET_PATTERN +
        r'(?:\s+' + APT_SUITE_UNIT_PATTERN + r')?' +
        US_CITY_ZIP_PATTERN
    )

    # Full Canadian Address pattern
    CA_ADDRESS_PATTERN = (
        US_STREET_PATTERN +
        r'(?:\s+' + APT_SUITE_UNIT_PATTERN + r')?' +
        CA_CITY_POSTAL_PATTERN
    )

    # Combined address pattern (US or Canadian)
    ADDRESS_PATTERN = r'(?:' + US_ADDRESS_PATTERN + r'|' + CA_ADDRESS_PATTERN + r')'
    
    @classmethod
    def redact(cls, text: str) -> str:
        """
        Redact all confidential information from text.

        Order matters: addresses are redacted before SSN to prevent
        ZIP+4 codes (e.g., 78701-1234) from being matched as SSN.

        Args:
            text: Raw parsed text from resume

        Returns:
            Text with PII redacted
        """
        # Redact emails first (most distinct pattern)
        text = cls._redact_emails(text)

        # Redact phone numbers
        text = cls._redact_phones(text)

        # Redact addresses BEFORE SSN (to handle ZIP+4 correctly)
        text = cls._redact_addresses(text)

        # Redact SSNs (after addresses to avoid ZIP+4 conflicts)
        text = cls._redact_ssn(text)

        # Redact dates of birth
        text = cls._redact_dates_of_birth(text)

        return text
    
    @classmethod
    def _redact_emails(cls, text: str) -> str:
        """Redact email addresses."""
        return re.sub(cls.EMAIL_PATTERN, '[EMAIL_REDACTED]', text, flags=re.IGNORECASE)
    
    @classmethod
    def _redact_phones(cls, text: str) -> str:
        """
        Redact phone numbers using regex pattern with phonenumbers library validation.

        Supports:
        - US local formats: (123) 456-7890, 123-456-7890, 123.456.7890
        - International formats: +44 20 7946 0958, +91 98765 43210
        - Extensions: x123, ext. 456, extension 789 (captured with phone)
        - Alphanumeric: 1-800-FLOWERS
        - Toll-free: 1-800-555-1234

        Args:
            text: Raw text containing phone numbers

        Returns:
            Text with phone numbers redacted (including extensions)
        """
        def replace_phone(match):
            """Validate and replace phone numbers."""
            phone_str = match.group(0)

            # For alphanumeric numbers (1-800-FLOWERS), redact directly
            if re.search(r'[A-Za-z]', phone_str):
                return '[PHONE_REDACTED]'

            # Validate with phonenumbers library
            try:
                # Try to parse with default region US
                parsed = phonenumbers.parse(phone_str, "US")
                # Use is_possible_number to catch phone-like patterns
                if phonenumbers.is_possible_number(parsed):
                    return '[PHONE_REDACTED]'
            except phonenumbers.NumberParseException:
                pass

            # Try parsing as international format (add + if missing)
            if not phone_str.startswith('+'):
                try:
                    parsed = phonenumbers.parse('+' + phone_str.replace('+', ''), "US")
                    if phonenumbers.is_possible_number(parsed):
                        return '[PHONE_REDACTED]'
                except phonenumbers.NumberParseException:
                    pass

            # If validation fails, return original string
            return phone_str

        return re.sub(cls.PHONE_PATTERN, replace_phone, text, flags=re.IGNORECASE)
    
    @classmethod
    def _redact_ssn(cls, text: str) -> str:
        """Redact Social Security Numbers."""
        return re.sub(cls.SSN_PATTERN, '[SSN_REDACTED]', text)
    
    @classmethod
    def _redact_dates_of_birth(cls, text: str) -> str:
        """Redact dates of birth."""
        return re.sub(cls.DOB_PATTERN, '[DOB_REDACTED]', text, flags=re.IGNORECASE)
    
    @classmethod
    def _redact_addresses(cls, text: str) -> str:
        """
        Redact street addresses from text.
        
        Supports:
        - US addresses: 123 Main St, Springfield, IL 62701
        - US with apartment: 456 Oak Ave Apt 7B, Chicago, IL 60601
        - US with suite: 789 Business Blvd Suite 100, NY, NY 10001
        - US ZIP+4: 321 Elm Dr, Austin, TX 78701-1234
        - Canadian addresses: 123 Maple St, Toronto, ON M5V 2T6
        - PO Boxes: PO Box 1234, Seattle, WA 98101
        - Rural Routes: RR 1 Box 234, Nashville, TN 37201
        
        Args:
            text: Raw text containing addresses
            
        Returns:
            Text with addresses redacted
        """
        # Redact PO Boxes first (more specific pattern)
        text = re.sub(cls.PO_BOX_PATTERN, '[ADDRESS_REDACTED]', text, flags=re.IGNORECASE)
        
        # Redact Rural Routes
        text = re.sub(cls.RURAL_ROUTE_PATTERN, '[ADDRESS_REDACTED]', text, flags=re.IGNORECASE)
        
        # Redact full US and Canadian addresses
        text = re.sub(cls.ADDRESS_PATTERN, '[ADDRESS_REDACTED]', text, flags=re.IGNORECASE)
        
        return text
