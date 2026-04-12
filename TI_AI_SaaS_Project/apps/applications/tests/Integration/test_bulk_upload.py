"""
Integration Tests for Bulk Upload Workflow

Tests the complete bulk upload workflow including:
- Batch initialization
- File upload
- Duplicate detection
- Batch commit

Note: Celery tasks are run synchronously during testing via CELERY_TASK_ALWAYS_EAGER.
"""

from django.test import TestCase, Client, override_settings
from django.contrib.auth import get_user_model
from django.core.files.uploadedfile import SimpleUploadedFile
from django.core.cache import cache
from django.urls import reverse
from apps.jobs.models import JobListing
from apps.applications.models import UploadBatch, Applicant
from apps.accounts.models import UserProfile
from apps.applications.services.resume_parser import ResumeParserService
from apps.applications.tasks import finalize_bulk_upload_batch
from datetime import timedelta
from django.utils import timezone
import zipfile
import json
from io import BytesIO
from docx import Document

User = get_user_model()


# Configure Celery to run tasks synchronously during testing
@override_settings(CELERY_TASK_ALWAYS_EAGER=True)
class BulkUploadWorkflowIntegrationTest(TestCase):
    """Integration tests for the complete bulk upload workflow."""

    def setUp(self):
        """Set up test data for each test."""
        # Clear cache to avoid rate limiting
        cache.clear()

        self.client = Client()

        # Create test user (TAS)
        self.user = User.objects.create_user(
            username='testuser_bulk',
            email='test_bulk@example.com',
            password='testpass123'
        )

        # Create user profile (required by RBAC middleware)
        UserProfile.objects.create(
            user=self.user,
            is_talent_acquisition_specialist=True
        )
        
        # Login to get JWT cookies (using the actual login endpoint)
        login_response = self.client.post(
            reverse('api:login'),
            data=json.dumps({
                'username': 'testuser_bulk',
                'password': 'testpass123'
            }),
            content_type='application/json'
        )
        
        # Verify login was successful
        self.assertEqual(login_response.status_code, 200)
        self.assertIn('access_token', self.client.cookies)
        
        # Create job listing
        self.job_listing = JobListing.objects.create(
            title='Test Job',
            description='Test Description',
            required_skills=['Python'],
            required_experience=2,
            job_level='Junior',
            start_date=timezone.now() + timedelta(days=1),
            expiration_date=timezone.now() + timedelta(days=30),
            upload_type='bulk',
            status='Active',
            created_by=self.user
        )

        # Create a valid DOCX file for testing (50KB minimum)
        self.docx_content = self._create_valid_docx()

    def _create_valid_docx(self, unique_id='default'):
        """Create a valid DOCX file (50KB minimum as per requirements).

        Args:
            unique_id: A unique identifier to include in the document to ensure
                      each generated file has unique content (for avoiding duplicate detection).
                      Should be an integer for best results.
        """
        # Generate unique phone number based on unique_id
        # Use valid US phone number format that passes phonenumbers validation
        # Area code 202 (Washington DC) with unique last 4 digits
        try:
            uid = int(unique_id)
        except (ValueError, TypeError):
            uid = 0

        # Generate unique last 4 digits (0000-9999)
        last_four = uid % 10000
        # Format: (202) 555-XXXX where XXXX is unique
        phone_suffix = f"(202) 555-{last_four:04d}"

        # Create a real DOCX document using python-docx with basic content
        doc = Document()
        doc.add_paragraph(f'John Doe {unique_id}')
        doc.add_paragraph('Software Engineer')
        doc.add_paragraph(f'Email: john.doe.{unique_id}@example.com')
        doc.add_paragraph(f'Phone: {phone_suffix}')
        doc.add_paragraph('Experience: 5 years in Python development')

        # Add substantial content to make it a realistic resume
        # Each paragraph adds roughly 200-300 bytes, so we need ~200 paragraphs for 50KB+
        # AVOID including patterns that could be mistakenly extracted as phone numbers:
        # - No standalone sequences of 10+ digits
        # - No patterns like (XXX) XXX-XXXX in the content
        for i in range(200):
            doc.add_paragraph(f'Work item {unique_id}-{i}: Developed software solutions using Python, Django, REST APIs, Flask, FastAPI, PostgreSQL, MongoDB, Redis, Docker, Kubernetes, cloud platforms, CI/CD pipelines, Jenkins, GitLab, GitHub Actions, Terraform, Ansible, microservices, event-driven architecture, message queues, RabbitMQ, Kafka, Elasticsearch, Kibana, Prometheus, Grafana, and various modern web technologies to deliver high-quality products.')

        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        docx_bytes = buffer.getvalue()

        # DOCX files are ZIP archives - we can add padding via ZIP comment
        # The ZIP comment is stored after the end of central directory record
        min_size = 51 * 1024  # 51KB to be safely above 50KB minimum

        if len(docx_bytes) < min_size:
            # Find the end of central directory signature (PK\x05\x06)
            eocd_sig = b'PK\x05\x06'
            eocd_pos = docx_bytes.rfind(eocd_sig)

            if eocd_pos != -1:
                # Comment length is at offset 20 from EOCD signature
                comment_len_pos = eocd_pos + 20

                # Create new comment with padding (use safe ASCII characters only)
                padding_needed = min_size - len(docx_bytes) + 100  # Extra buffer
                new_comment = b'A' * min(padding_needed, 65000)  # ZIP comment max is 65535

                # Update comment length in little-endian format
                docx_bytes = bytearray(docx_bytes)
                docx_bytes[comment_len_pos:comment_len_pos+2] = len(new_comment).to_bytes(2, 'little')

                # Truncate at EOCD end and append new comment
                docx_bytes = bytes(docx_bytes[:eocd_pos + 22]) + new_comment

        docx_bytes = bytes(docx_bytes)

        # Verify minimum size (50KB = 51200 bytes)
        assert len(docx_bytes) >= 50 * 1024, f"DOCX file too small: {len(docx_bytes)} bytes, need {50 * 1024}"

        # Verify it's still a valid ZIP/DOCX file
        try:
            with zipfile.ZipFile(BytesIO(docx_bytes), 'r') as zf:
                zf.namelist()
        except zipfile.BadZipFile as e:
            raise AssertionError(f"Generated invalid DOCX file: {e}")

        return docx_bytes

    def test_complete_bulk_upload_workflow(self):
        """Test complete bulk upload workflow from init to commit."""
        # Step 1: Initialize batch
        init_response = self.client.post(
            '/api/applications/bulk-upload/init/',
            content_type='application/json',
            data=json.dumps({'job_listing_id': str(self.job_listing.id)})
        )

        self.assertEqual(init_response.status_code, 201)
        init_data = init_response.json()
        self.assertIn('batch_id', init_data)
        self.assertEqual(init_data['batch_number'], 1)

        batch_id = init_data['batch_id']

        # Step 2: Upload files with unique content
        uploaded_files = []
        for i in range(5):
            # Create unique content for each file to avoid duplicate detection
            # Use numeric unique_id to ensure unique phone numbers
            unique_docx = self._create_valid_docx(unique_id=i)
            upload_file = SimpleUploadedFile(
                f"test_resume_{i}.docx",
                unique_docx,
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

            upload_response = self.client.post(
                '/api/applications/bulk-upload/upload/',
                {
                    'batch_id': batch_id,
                    'file': upload_file
                }
            )
            self.assertEqual(upload_response.status_code, 200)
            upload_data = upload_response.json()
            self.assertIn('file_id', upload_data)
            uploaded_files.append(upload_data)

        # Step 3: Validate batch
        validate_response = self.client.post(
            '/api/applications/bulk-upload/validate/',
            content_type='application/json',
            data=json.dumps({'batch_id': batch_id})
        )
        self.assertEqual(validate_response.status_code, 200)
        validate_data = validate_response.json()
        self.assertEqual(validate_data['total_files'], 5)
        self.assertEqual(validate_data['status'], 'awaiting_review')

        # Step 4: Commit batch (triggers async processing)
        # With CELERY_TASK_ALWAYS_EAGER, tasks run synchronously
        commit_response = self.client.post(
            '/api/applications/bulk-upload/commit/',
            content_type='application/json',
            data=json.dumps({'batch_id': batch_id})
        )
        self.assertEqual(commit_response.status_code, 202)  # Accepted for async processing
        commit_data = commit_response.json()
        self.assertEqual(commit_data['status'], 'processing')
        self.assertEqual(commit_data['total_files'], 5)

        # Note: With CELERY_TASK_ALWAYS_EAGER, the process_bulk_upload_batch task
        # runs synchronously and dispatches process_resume_async for each file.
        # However, finalize_bulk_upload_batch needs to be called separately to
        # update JobListing counters and set batch status to 'committed'.
        # In production, this would be triggered when all file tasks complete.
        
        # Refresh batch to get latest progress
        batch = UploadBatch.objects.get(id=batch_id)

        # Finalize the batch to update status (pass empty results list for chord compatibility)
        finalize_bulk_upload_batch([], batch_id)

        # Refresh batch again after finalize
        batch.refresh_from_db()

        # Verify JobListing updated
        self.job_listing.refresh_from_db()
        # batch_count should NOT increment because only 5 files were committed (not 100)
        self.assertEqual(self.job_listing.batch_count, 0)
        # total_resumes should reflect the 5 committed files
        self.assertEqual(self.job_listing.total_resumes, 5)

        # Verify Applicants created
        applicants_count = Applicant.objects.filter(
            job_listing=self.job_listing,
            upload_batch_id=batch_id
        ).count()
        self.assertEqual(applicants_count, 5)

    def test_bulk_upload_with_duplicate_skip(self):
        """Test bulk upload workflow with duplicate detection and skip."""
        # Create a file that will be used for both existing applicant and upload
        duplicate_docx = self._create_valid_docx(unique_id=5000)
        
        # Calculate the file hash that will be generated
        file_hash = ResumeParserService.calculate_file_hash(duplicate_docx)

        # Create existing applicant with the SAME resume hash
        existing_applicant = Applicant.objects.create(
            job_listing=self.job_listing,
            first_name='Existing',
            last_name='User',
            email='existing@example.com',
            phone='+1234567890',
            resume_file_hash=file_hash,
            resume_parsed_text='Test resume'
        )

        # Initialize batch
        init_response = self.client.post(
            '/api/applications/bulk-upload/init/',
            content_type='application/json',
            data=json.dumps({'job_listing_id': str(self.job_listing.id)})
        )
        batch_id = init_response.json()['batch_id']

        # Upload the SAME file (will be detected as duplicate)
        upload_file = SimpleUploadedFile(
            "test_resume_duplicate.docx",
            duplicate_docx,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

        upload_response = self.client.post(
            '/api/applications/bulk-upload/upload/',
            {
                'batch_id': batch_id,
                'file': upload_file
            }
        )
        self.assertEqual(upload_response.status_code, 200)

        # Validate batch - this will detect the duplicate
        validate_response = self.client.post(
            '/api/applications/bulk-upload/validate/',
            content_type='application/json',
            data=json.dumps({'batch_id': batch_id})
        )
        self.assertEqual(validate_response.status_code, 200)
        validate_data = validate_response.json()
        
        # Verify duplicate was detected
        self.assertEqual(validate_data['status'], 'awaiting_review')
        self.assertGreater(len(validate_data.get('duplicates', [])), 0)

        # Submit decision to skip all duplicates
        decisions_response = self.client.post(
            '/api/applications/bulk-upload/decisions/',
            content_type='application/json',
            data=json.dumps({
                'batch_id': batch_id,
                'decisions': [{'action': 'skip_all'}]
            })
        )
        self.assertEqual(decisions_response.status_code, 200)

        # Commit batch (triggers async processing)
        # With CELERY_TASK_ALWAYS_EAGER, tasks run synchronously
        commit_response = self.client.post(
            '/api/applications/bulk-upload/commit/',
            content_type='application/json',
            data=json.dumps({'batch_id': batch_id})
        )
        self.assertEqual(commit_response.status_code, 202)  # Accepted for async processing

        # Finalize the batch to update status (pass empty results list for chord compatibility)
        finalize_bulk_upload_batch([], batch_id)

        # Verify no new applicants created (duplicate was skipped)
        applicants_count = Applicant.objects.filter(
            job_listing=self.job_listing,
            upload_batch_id=batch_id
        ).count()
        self.assertEqual(applicants_count, 0)

    def test_bulk_upload_cancel_cleanup(self):
        """Test that cancelling a batch cleans up temporary files."""
        # Initialize batch
        init_response = self.client.post(
            '/api/applications/bulk-upload/init/',
            content_type='application/json',
            data=json.dumps({'job_listing_id': str(self.job_listing.id)})
        )
        batch_id = init_response.json()['batch_id']

        # Upload a file with unique content
        upload_file = SimpleUploadedFile(
            "test_resume.docx",
            self._create_valid_docx(unique_id=2000),  # Unique ID for cancel test
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

        upload_response = self.client.post(
            '/api/applications/bulk-upload/upload/',
            {
                'batch_id': batch_id,
                'file': upload_file
            }
        )
        self.assertEqual(upload_response.status_code, 200)

        # Cancel batch
        cancel_response = self.client.delete(
            f'/api/applications/bulk-upload/cancel/{batch_id}/'
        )
        self.assertEqual(cancel_response.status_code, 200)

        # Verify batch status
        batch = UploadBatch.objects.get(id=batch_id)
        self.assertEqual(batch.status, 'cancelled')

    def test_bulk_upload_limits_enforcement(self):
        """Test that upload limits are enforced."""
        # Initialize batch
        init_response = self.client.post(
            '/api/applications/bulk-upload/init/',
            content_type='application/json',
            data=json.dumps({'job_listing_id': str(self.job_listing.id)})
        )
        batch_id = init_response.json()['batch_id']

        # Try to upload 101 files (should fail at 100)
        for i in range(101):
            # Create unique content for each file to avoid duplicate detection
            # Use offset of 10000 to avoid collision with other tests
            unique_docx = self._create_valid_docx(unique_id=10000 + i)
            upload_file = SimpleUploadedFile(
                f"test_resume_{i}.docx",
                unique_docx,
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

            upload_response = self.client.post(
                '/api/applications/bulk-upload/upload/',
                {
                    'batch_id': batch_id,
                    'file': upload_file
                }
            )

            if i < 100:
                self.assertEqual(upload_response.status_code, 200)
            else:
                # 101st file should be rejected
                self.assertEqual(upload_response.status_code, 400)
                # DRF returns validation errors in field-specific format
                response_data = upload_response.json()
                error_message = ''
                if 'batch_id' in response_data:
                    error_message = ' '.join(response_data['batch_id']) if isinstance(response_data['batch_id'], list) else str(response_data['batch_id'])
                elif 'error' in response_data:
                    error_message = response_data['error']
                elif 'non_field_errors' in response_data:
                    error_message = ' '.join(response_data['non_field_errors']) if isinstance(response_data['non_field_errors'], list) else str(response_data['non_field_errors'])
                
                self.assertIn('Batch is full', error_message)

    def test_bulk_upload_form_type_rejected(self):
        """Test that bulk upload is rejected for form-type job listings."""
        form_job = JobListing.objects.create(
            title='Form Job',
            description='Form Description',
            required_skills=['Python'],
            required_experience=2,
            job_level='Junior',
            start_date=timezone.now() + timedelta(days=1),
            expiration_date=timezone.now() + timedelta(days=30),
            upload_type='form',
            status='Active',
            created_by=self.user
        )

        init_response = self.client.post(
            '/api/applications/bulk-upload/init/',
            content_type='application/json',
            data=json.dumps({'job_listing_id': str(form_job.id)})
        )

        self.assertEqual(init_response.status_code, 400)
        # DRF returns validation errors in field-specific format
        response_data = init_response.json()
        error_message = ''
        if 'job_listing_id' in response_data:
            error_message = ' '.join(response_data['job_listing_id']) if isinstance(response_data['job_listing_id'], list) else str(response_data['job_listing_id'])
        elif 'error' in response_data:
            error_message = response_data['error']

        self.assertIn('bulk upload', error_message.lower())
