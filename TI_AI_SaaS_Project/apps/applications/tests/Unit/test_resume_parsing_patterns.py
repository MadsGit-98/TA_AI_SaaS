"""
Tests for Resume Parsing Service SSN and DOB Patterns

Tests for confidential information filtering patterns.
"""

from django.test import SimpleTestCase
from services.resume_parsing_service import ResumeParserService, ConfidentialInfoFilter
from io import BytesIO
from docx import Document
from docx.shared import Inches


class ResumeParserServiceDocxTableTests(SimpleTestCase):
    """Tests for DOCX table text extraction."""

    def test_extract_text_from_docx_with_tables(self):
        """Test that text inside tables is extracted."""
        # Create a DOCX with a table
        doc = Document()
        doc.add_paragraph("Header text")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Skill 1"
        table.cell(0, 1).text = "Expert"
        table.cell(1, 0).text = "Skill 2"
        table.cell(1, 1).text = "Advanced"
        doc.add_paragraph("Footer text")

        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Extract text
        result = ResumeParserService.extract_text_from_docx(buffer.getvalue())

        # Verify table content is included
        self.assertIn("Skill 1", result)
        self.assertIn("Expert", result)
        self.assertIn("Skill 2", result)
        self.assertIn("Advanced", result)
        self.assertIn("Header text", result)
        self.assertIn("Footer text", result)

    def test_extract_text_from_docx_table_avoids_duplicates(self):
        """Test that duplicate table cell text is not added multiple times."""
        doc = Document()
        table = doc.add_table(rows=3, cols=1)
        table.cell(0, 0).text = "Same text"
        table.cell(1, 0).text = "Same text"  # Duplicate
        table.cell(2, 0).text = "Different"

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        result = ResumeParserService.extract_text_from_docx(buffer.getvalue())

        # "Same text" should appear only once
        self.assertEqual(result.count("Same text"), 1)
        self.assertIn("Different", result)

    def test_extract_text_from_docx_paragraphs_only(self):
        """Test extraction from DOCX without tables (paragraphs only)."""
        doc = Document()
        doc.add_paragraph("Paragraph 1")
        doc.add_paragraph("Paragraph 2")

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        result = ResumeParserService.extract_text_from_docx(buffer.getvalue())

        self.assertIn("Paragraph 1", result)
        self.assertIn("Paragraph 2", result)


class ConfidentialInfoFilterSSNTests(SimpleTestCase):
    """Tests for SSN pattern matching and redaction."""

    def test_redact_ssn_formatted_with_dashes(self):
        """Test SSN in XXX-XX-XXXX format."""
        text = "My SSN is 123-45-6789 for verification."
        result = ConfidentialInfoFilter.redact(text)
        self.assertIn("[SSN_REDACTED]", result)
        self.assertNotIn("123-45-6789", result)

    def test_redact_ssn_with_spaces(self):
        """Test SSN in XXX XX XXXX format."""
        text = "My SSN is 123 45 6789 for verification."
        result = ConfidentialInfoFilter.redact(text)
        self.assertIn("[SSN_REDACTED]", result)
        self.assertNotIn("123 45 6789", result)

    def test_redact_ssn_no_separator(self):
        """Test SSN in XXXXXXXXXX format (9 consecutive digits)."""
        text = "My SSN is 123456789 for verification."
        result = ConfidentialInfoFilter.redact(text)
        self.assertIn("[SSN_REDACTED]", result)
        self.assertNotIn("123456789", result)

    def test_redact_multiple_ssns(self):
        """Test multiple SSNs in different formats."""
        text = "SSN1: 111-22-3333, SSN2: 444 55 6666, SSN3: 777889999"
        result = ConfidentialInfoFilter.redact(text)
        self.assertEqual(result.count("[SSN_REDACTED]"), 3)

    def test_ssn_word_boundary_start(self):
        """Test that SSN pattern respects word boundary at start."""
        text = "Invalid: 1123-45-6789 should not match."
        result = ConfidentialInfoFilter.redact(text)
        # Should not redact because it has extra digit at start
        self.assertIn("1123-45-6789", result)

    def test_ssn_word_boundary_end(self):
        """Test that SSN pattern respects word boundary at end."""
        text = "Invalid: 123-45-67890 should not match."
        result = ConfidentialInfoFilter.redact(text)
        # Should not redact because it has extra digit at end
        self.assertIn("123-45-67890", result)


class ConfidentialInfoFilterDOBTests(SimpleTestCase):
    """Tests for DOB pattern matching and redaction."""

    # US Format Tests (MM/DD/YYYY and MM-DD-YYYY)
    def test_redact_dob_us_format_slashes(self):
        """Test DOB in MM/DD/YYYY format."""
        text = "DOB: 01/15/1990"
        result = ConfidentialInfoFilter.redact(text)
        self.assertIn("[DOB_REDACTED]", result)
        self.assertNotIn("01/15/1990", result)

    def test_redact_dob_us_format_dashes(self):
        """Test DOB in MM-DD-YYYY format."""
        text = "DOB: 12-25-1985"
        result = ConfidentialInfoFilter.redact(text)
        self.assertIn("[DOB_REDACTED]", result)
        self.assertNotIn("12-25-1985", result)

    # ISO Format Tests (YYYY-MM-DD)
    def test_redact_dob_iso_format(self):
        """Test DOB in ISO format YYYY-MM-DD."""
        text = "DOB: 1990-01-15"
        result = ConfidentialInfoFilter.redact(text)
        self.assertIn("[DOB_REDACTED]", result)
        self.assertNotIn("1990-01-15", result)

    def test_redact_dob_iso_format_2000s(self):
        """Test DOB in ISO format with 2000s year."""
        text = "DOB: 2005-06-30"
        result = ConfidentialInfoFilter.redact(text)
        self.assertIn("[DOB_REDACTED]", result)
        self.assertNotIn("2005-06-30", result)

    # Written Month Format Tests (Month DD, YYYY)
    def test_redact_dob_written_month_long(self):
        """Test DOB with written month name (long form)."""
        text = "DOB: January 15, 1990"
        result = ConfidentialInfoFilter.redact(text)
        self.assertIn("[DOB_REDACTED]", result)
        self.assertNotIn("January 15, 1990", result)

    def test_redact_dob_written_month_no_comma(self):
        """Test DOB with written month name without comma."""
        text = "DOB: January 15 1990"
        result = ConfidentialInfoFilter.redact(text)
        self.assertIn("[DOB_REDACTED]", result)
        self.assertNotIn("January 15 1990", result)

    def test_redact_dob_written_month_lowercase(self):
        """Test DOB with lowercase month name (case insensitive)."""
        text = "DOB: january 15, 1990"
        result = ConfidentialInfoFilter.redact(text)
        self.assertIn("[DOB_REDACTED]", result)
        self.assertNotIn("january 15, 1990", result)

    def test_redact_dob_written_month_all_months(self):
        """Test DOB with all month names."""
        months = [
            "January", "February", "March", "April", "May", "June",
            "July", "August", "September", "October", "November", "December"
        ]
        for month in months:
            text = f"DOB: {month} 15, 1990"
            result = ConfidentialInfoFilter.redact(text)
            self.assertIn("[DOB_REDACTED]", result, f"Failed for month: {month}")

    # Day-First Written Format Tests (DD Month YYYY)
    def test_redact_dob_day_first_written(self):
        """Test DOB in DD Month YYYY format."""
        text = "DOB: 15 January 1990"
        result = ConfidentialInfoFilter.redact(text)
        self.assertIn("[DOB_REDACTED]", result)
        self.assertNotIn("15 January 1990", result)

    def test_redact_dob_day_first_written_lowercase(self):
        """Test DOB in DD Month YYYY format with lowercase month."""
        text = "DOB: 15 january 1990"
        result = ConfidentialInfoFilter.redact(text)
        self.assertIn("[DOB_REDACTED]", result)
        self.assertNotIn("15 january 1990", result)

    def test_redact_dob_day_first_with_comma(self):
        """Test DOB in DD Month, YYYY format."""
        text = "DOB: 15 January, 1990"
        result = ConfidentialInfoFilter.redact(text)
        self.assertIn("[DOB_REDACTED]", result)
        self.assertNotIn("15 January, 1990", result)

    # Edge Cases and False Positive Prevention
    def test_dob_year_range_1900(self):
        """Test DOB with year 1900 (valid range start)."""
        text = "DOB: 01/15/1900"
        result = ConfidentialInfoFilter.redact(text)
        self.assertIn("[DOB_REDACTED]", result)

    def test_dob_year_range_2099(self):
        """Test DOB with year 2099 (valid range end)."""
        text = "DOB: 01/15/2099"
        result = ConfidentialInfoFilter.redact(text)
        self.assertIn("[DOB_REDACTED]", result)

    def test_dob_year_out_of_range_1899(self):
        """Test that year 1899 does not match (out of range)."""
        text = "DOB: 01/15/1899"
        result = ConfidentialInfoFilter.redact(text)
        self.assertIn("1899", result)

    def test_dob_year_out_of_range_2100(self):
        """Test that year 2100 does not match (out of range)."""
        text = "DOB: 01/15/2100"
        result = ConfidentialInfoFilter.redact(text)
        self.assertIn("2100", result)

    def test_dob_word_boundary_prevents_partial_match(self):
        """Test that word boundaries prevent partial matches."""
        text = "Not a DOB: 12345/12/1990 should not match."
        result = ConfidentialInfoFilter.redact(text)
        self.assertIn("12345", result)

    def test_redact_multiple_dobs_different_formats(self):
        """Test multiple DOBs in different formats."""
        text = "Person1: 01/15/1990, Person2: 1985-06-20, Person3: March 10, 1988"
        result = ConfidentialInfoFilter.redact(text)
        self.assertEqual(result.count("[DOB_REDACTED]"), 3)
