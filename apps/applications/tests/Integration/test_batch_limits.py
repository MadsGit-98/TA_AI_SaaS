"""
Integration Tests for Batch Upload Limits Enforcement

Tests that batch upload limits are properly enforced.

Note: Celery tasks are run synchronously during testing via CELERY_TASK_ALWAYS_EAGER.
"""

from django.test import TestCase, Client, override_settings
from django.core.files.uploadedfile import SimpleUploadedFile
from django.core.cache import cache
from django.urls import reverse
from django.contrib.auth import get_user_model
from apps.jobs.models import JobListing
from apps.applications.models import UploadBatch
from apps.accounts.models import UserProfile
from datetime import timedelta
from django.utils import timezone
import json
from io import BytesIO
from docx import Document
from apps.applications.tasks import finalize_bulk_upload_batch

User = get_user_model()


@override_settings(CELERY_TASK_ALWAYS_EAGER=True)
class BatchLimitsIntegrationTest(TestCase):
    """Integration tests for batch upload limits enforcement."""

    def setUp(self):
        """Set up test data for each test."""
        # Clear cache to avoid rate limiting
        cache.clear()
        
        self.client = Client()
        
        # Create test user (TAS)
        self.user = User.objects.create_user(
            username='testuser_limits',
            email='test_limits@example.com',
            password='testpass123'
        )
        
        # Create user profile (required by RBAC middleware)
        UserProfile.objects.create(
            user=self.user,
            is_talent_acquisition_specialist=True
        )
        
        # Login to get JWT cookies (using the actual login endpoint)
        login_response = self.client.post(
            reverse('api:login'),
            data=json.dumps({
                'username': 'testuser_limits',
                'password': 'testpass123'
            }),
            content_type='application/json'
        )
        
        # Verify login was successful
        self.assertEqual(login_response.status_code, 200)
        self.assertIn('access_token', self.client.cookies)
        
        # Create job listing
        self.job_listing = JobListing.objects.create(
            title='Test Job',
            description='Test Description',
            required_skills=['Python'],
            required_experience=2,
            job_level='Junior',
            start_date=timezone.now() + timedelta(days=1),
            expiration_date=timezone.now() + timedelta(days=30),
            upload_type='bulk',
            created_by=self.user
        )

        # Create a valid DOCX file for testing (50KB minimum)
        self.docx_content = self._create_valid_docx()

    def _create_valid_docx(self, unique_id='default'):
        """Create a valid DOCX file (50KB minimum as per requirements).

        Args:
            unique_id: A unique identifier to include in the document to ensure
                      each generated file has unique content (for avoiding duplicate detection).
                      Should be an integer for best results.
        """
        import zipfile

        # Generate unique phone number based on unique_id
        # Use valid US phone number format that passes phonenumbers validation
        # Area code 202 (Washington DC) with unique last 4 digits
        try:
            uid = int(unique_id)
        except (ValueError, TypeError):
            uid = 0

        # Generate unique last 4 digits (0000-9999)
        last_four = uid % 10000
        # Format: (202) 555-XXXX where XXXX is unique
        phone_suffix = f"(202) 555-{last_four:04d}"

        # Create a real DOCX document using python-docx with basic content
        doc = Document()
        doc.add_paragraph(f'Test User {unique_id}')
        doc.add_paragraph('Test Position')
        doc.add_paragraph(f'Email: test.{unique_id}@example.com')
        doc.add_paragraph(f'Phone: {phone_suffix}')
        doc.add_paragraph('Experience: 5 years in Python development')

        # Add substantial content to make it a realistic resume
        # Each paragraph adds roughly 200-300 bytes, so we need ~200 paragraphs for 50KB+
        for i in range(200):
            doc.add_paragraph(f'Work experience item {i} ({unique_id}): Developed software solutions using Python, Django, REST APIs, Flask, FastAPI, PostgreSQL, MongoDB, Redis, Docker, Kubernetes, AWS, Azure, GCP, CI/CD pipelines, Jenkins, GitLab, GitHub Actions, Terraform, Ansible, microservices, event-driven architecture, message queues, RabbitMQ, Kafka, Elasticsearch, Kibana, Prometheus, Grafana, and various modern technologies to deliver high-quality products for enterprise clients.')

        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        docx_bytes = buffer.getvalue()

        # DOCX files are ZIP archives - we can add padding via ZIP comment
        # The ZIP comment is stored after the end of central directory record
        min_size = 51 * 1024  # 51KB to be safely above 50KB minimum
        
        if len(docx_bytes) < min_size:
            # Find the end of central directory signature (PK\x05\x06)
            eocd_sig = b'PK\x05\x06'
            eocd_pos = docx_bytes.rfind(eocd_sig)
            
            if eocd_pos != -1:
                # Comment length is at offset 20 from EOCD signature
                comment_len_pos = eocd_pos + 20
                
                # Create new comment with padding (use safe ASCII characters only)
                padding_needed = min_size - len(docx_bytes) + 100  # Extra buffer
                new_comment = b'A' * min(padding_needed, 65000)  # ZIP comment max is 65535
                
                # Update comment length in little-endian format
                docx_bytes = bytearray(docx_bytes)
                docx_bytes[comment_len_pos:comment_len_pos+2] = len(new_comment).to_bytes(2, 'little')
                
                # Truncate at EOCD end and append new comment
                docx_bytes = bytes(docx_bytes[:eocd_pos + 22]) + new_comment

        docx_bytes = bytes(docx_bytes)
        
        # Verify minimum size (50KB = 51200 bytes)
        assert len(docx_bytes) >= 50 * 1024, f"DOCX file too small: {len(docx_bytes)} bytes, need {50 * 1024}"
        
        # Verify it's still a valid ZIP/DOCX file
        try:
            with zipfile.ZipFile(BytesIO(docx_bytes), 'r') as zf:
                zf.namelist()
        except zipfile.BadZipFile as e:
            raise AssertionError(f"Generated invalid DOCX file: {e}")

        return docx_bytes

    def test_batch_count_limit_3(self):
        """Test that batch_count increments only when 100 files are committed."""
        # The new logic: batch_count only increments when a full batch (100 files) is committed
        # Users can upload multiple times until total_resumes reaches 300
        
        # Create first batch and commit 100 files
        init_response = self.client.post(
            '/api/applications/bulk-upload/init/',
            content_type='application/json',
            data=json.dumps({'job_listing_id': str(self.job_listing.id)})
        )
        self.assertEqual(init_response.status_code, 201)
        batch_id = init_response.json()['batch_id']
        
        # Upload 100 files with unique content
        for i in range(100):
            unique_docx = self._create_valid_docx(unique_id=i)
            upload_file = SimpleUploadedFile(
                f"test_resume_{i}.docx",
                unique_docx,
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            upload_response = self.client.post(
                '/api/applications/bulk-upload/upload/',
                {
                    'batch_id': batch_id,
                    'file': upload_file
                }
            )
            self.assertEqual(upload_response.status_code, 200)
        
        # Validate and commit
        validate_response = self.client.post(
            '/api/applications/bulk-upload/validate/',
            content_type='application/json',
            data=json.dumps({'batch_id': batch_id})
        )
        self.assertEqual(validate_response.status_code, 200)
        
        commit_response = self.client.post(
            '/api/applications/bulk-upload/commit/',
            content_type='application/json',
            data=json.dumps({'batch_id': batch_id})
        )
        self.assertEqual(commit_response.status_code, 202)  # Accepted for async processing
        
        # Finalize the batch to complete processing
        finalize_bulk_upload_batch([], batch_id)

        # Refresh from database - batch_count should now be 1
        self.job_listing.refresh_from_db()
        self.assertEqual(self.job_listing.batch_count, 1)
        self.assertEqual(self.job_listing.total_resumes, 100)
        
        # Create second batch and commit 100 files
        init_response2 = self.client.post(
            '/api/applications/bulk-upload/init/',
            content_type='application/json',
            data=json.dumps({'job_listing_id': str(self.job_listing.id)})
        )
        self.assertEqual(init_response2.status_code, 201)
        batch_id2 = init_response2.json()['batch_id']
        
        # Upload 100 files with unique content
        for i in range(100, 200):
            unique_docx = self._create_valid_docx(unique_id=i)
            upload_file = SimpleUploadedFile(
                f"test_resume_{i}.docx",
                unique_docx,
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            upload_response = self.client.post(
                '/api/applications/bulk-upload/upload/',
                {
                    'batch_id': batch_id2,
                    'file': upload_file
                }
            )
            self.assertEqual(upload_response.status_code, 200)
        
        # Validate and commit
        validate_response2 = self.client.post(
            '/api/applications/bulk-upload/validate/',
            content_type='application/json',
            data=json.dumps({'batch_id': batch_id2})
        )
        self.assertEqual(validate_response2.status_code, 200)
        
        commit_response2 = self.client.post(
            '/api/applications/bulk-upload/commit/',
            content_type='application/json',
            data=json.dumps({'batch_id': batch_id2})
        )
        self.assertEqual(commit_response2.status_code, 202)  # Accepted for async processing
        
        # Finalize the batch
        finalize_bulk_upload_batch([], batch_id2)

        # Refresh from database - batch_count should now be 2
        self.job_listing.refresh_from_db()
        self.assertEqual(self.job_listing.batch_count, 2)
        self.assertEqual(self.job_listing.total_resumes, 200)
        
        # Create third batch and commit 100 files
        init_response3 = self.client.post(
            '/api/applications/bulk-upload/init/',
            content_type='application/json',
            data=json.dumps({'job_listing_id': str(self.job_listing.id)})
        )
        self.assertEqual(init_response3.status_code, 201)
        batch_id3 = init_response3.json()['batch_id']
        
        # Upload 100 files with unique content
        for i in range(200, 300):
            unique_docx = self._create_valid_docx(unique_id=i)
            upload_file = SimpleUploadedFile(
                f"test_resume_{i}.docx",
                unique_docx,
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            upload_response = self.client.post(
                '/api/applications/bulk-upload/upload/',
                {
                    'batch_id': batch_id3,
                    'file': upload_file
                }
            )
            self.assertEqual(upload_response.status_code, 200)
        
        # Validate and commit
        validate_response3 = self.client.post(
            '/api/applications/bulk-upload/validate/',
            content_type='application/json',
            data=json.dumps({'batch_id': batch_id3})
        )
        self.assertEqual(validate_response3.status_code, 200)
        
        commit_response3 = self.client.post(
            '/api/applications/bulk-upload/commit/',
            content_type='application/json',
            data=json.dumps({'batch_id': batch_id3})
        )
        self.assertEqual(commit_response3.status_code, 202)  # Accepted for async processing
        
        # Finalize the batch

        finalize_bulk_upload_batch([], batch_id3)

        # Refresh from database - batch_count should now be 3
        self.job_listing.refresh_from_db()
        self.assertEqual(self.job_listing.batch_count, 3)
        self.assertEqual(self.job_listing.total_resumes, 300)
        
        # Try to create another batch - should fail due to max resumes reached
        init_response4 = self.client.post(
            '/api/applications/bulk-upload/init/',
            content_type='application/json',
            data=json.dumps({'job_listing_id': str(self.job_listing.id)})
        )
        
        # Should fail because total_resumes is 300 (max limit)
        self.assertEqual(init_response4.status_code, 400)
        self.assertIn('Maximum resume limit reached', init_response4.json().get('error', ''))

    def test_total_resumes_limit_300(self):
        """Test that maximum 300 resumes per job listing is enforced."""
        # Set total resumes to 299
        self.job_listing.total_resumes = 299
        self.job_listing.save()

        init_response = self.client.post(
            '/api/applications/bulk-upload/init/',
            content_type='application/json',
            data=json.dumps({'job_listing_id': str(self.job_listing.id)})
        )

        # Should allow 1 more resume
        self.assertEqual(init_response.status_code, 201)
        self.assertEqual(init_response.json()['remaining_capacity'], 100)

        # Set total resumes to 300
        self.job_listing.total_resumes = 300
        self.job_listing.save()

        # Try to create another batch - should fail
        init_response = self.client.post(
            '/api/applications/bulk-upload/init/',
            content_type='application/json',
            data=json.dumps({'job_listing_id': str(self.job_listing.id)})
        )

        self.assertEqual(init_response.status_code, 400)
        self.assertIn('Maximum resume limit reached', init_response.json().get('error', ''))

    def test_file_count_limit_100_per_batch(self):
        """Test that maximum 100 files per batch is enforced."""
        # Initialize batch
        init_response = self.client.post(
            '/api/applications/bulk-upload/init/',
            content_type='application/json',
            data=json.dumps({'job_listing_id': str(self.job_listing.id)})
        )
        batch_id = init_response.json()['batch_id']

        # Upload 100 files with unique content
        for i in range(100):
            # Use numeric ID with offset to ensure unique phone numbers
            unique_docx = self._create_valid_docx(unique_id=100 + i)
            upload_file = SimpleUploadedFile(
                f"test_resume_{i}.docx",
                unique_docx,
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

            upload_response = self.client.post(
                '/api/applications/bulk-upload/upload/',
                {
                    'batch_id': batch_id,
                    'file': upload_file
                }
            )
            self.assertEqual(upload_response.status_code, 200, f"Failed at file {i+1}")

        # Try to upload 101st file - should fail
        upload_file = SimpleUploadedFile(
            "test_resume_101.docx",
            self._create_valid_docx(unique_id=201),  # Unique ID for 101st file
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

        upload_response = self.client.post(
            '/api/applications/bulk-upload/upload/',
            {
                'batch_id': batch_id,
                'file': upload_file
            }
        )

        self.assertEqual(upload_response.status_code, 400)
        # DRF returns validation errors in field-specific format
        # Check for the error message in batch_id field or as non-field error
        response_data = upload_response.json()
        error_message = ''
        if 'batch_id' in response_data:
            error_message = ' '.join(response_data['batch_id']) if isinstance(response_data['batch_id'], list) else str(response_data['batch_id'])
        elif 'error' in response_data:
            error_message = response_data['error']
        elif 'non_field_errors' in response_data:
            error_message = ' '.join(response_data['non_field_errors']) if isinstance(response_data['non_field_errors'], list) else str(response_data['non_field_errors'])
        
        self.assertIn('Batch is full', error_message)

    def test_remaining_capacity_calculation(self):
        """Test that remaining capacity is correctly calculated."""
        # Initialize batch
        init_response = self.client.post(
            '/api/applications/bulk-upload/init/',
            content_type='application/json',
            data=json.dumps({'job_listing_id': str(self.job_listing.id)})
        )

        self.assertEqual(init_response.json()['remaining_capacity'], 100)

        # Upload 25 files with unique content
        for i in range(25):
            # Use numeric ID with offset to ensure unique phone numbers
            unique_docx = self._create_valid_docx(unique_id=300 + i)
            upload_file = SimpleUploadedFile(
                f"test_resume_{i}.docx",
                unique_docx,
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

            self.client.post(
                '/api/applications/bulk-upload/upload/',
                {
                    'batch_id': init_response.json()['batch_id'],
                    'file': upload_file
                }
            )

        # Get status
        status_response = self.client.get(
            f'/api/applications/bulk-upload/status/{init_response.json()["batch_id"]}/'
        )

        self.assertEqual(status_response.json()['progress']['files_uploaded'], 25)

    def test_batch_number_no_longer_has_constraint(self):
        """Test that batch_number no longer has a max 3 constraint."""
        # The batch_number_max_3 constraint has been removed to allow
        # users to upload multiple times until 300 resumes is reached.
        # For example: 50 files x 6 batches = 300 files (6 batches total)
        
        # Create 5 batches with batch_number 1-5
        for i in range(1, 6):
            batch = UploadBatch.objects.create(
                job_listing=self.job_listing,
                batch_number=i,
                uploaded_by=self.user
            )
            self.assertEqual(batch.batch_number, i)
        
        # Creating a 6th batch should work (no constraint)
        batch_6 = UploadBatch.objects.create(
            job_listing=self.job_listing,
            batch_number=6,
            uploaded_by=self.user
        )
        self.assertEqual(batch_6.batch_number, 6)

    def test_file_count_constraint_database_level(self):
        """Test that file_count is properly tracked.
        
        Note: The file_count_max_100 database constraint was removed.
        File count limits are now enforced at the application level
        (API validation prevents uploading more than 100 files per batch).
        """
        batch = UploadBatch.objects.create(
            job_listing=self.job_listing,
            batch_number=1,
            uploaded_by=self.user
        )

        # file_count is automatically updated when files are added via add_file()
        # The 100 file limit is enforced at the API level, not database level
        self.assertEqual(batch.file_count, 0)
        
        # Verify file_count updates when files are added
        batch.add_file({
            'file_id': 'test-file-1',
            'filename': 'test.pdf',
            'file_hash': 'abc123',
            'size': 1000,
            'status': 'uploaded'
        })
        
        self.assertEqual(batch.file_count, 1)

    def test_upload_limits_reflected_in_job_listing(self):
        """Test that upload limits are properly reflected in JobListing."""
        # Initialize and commit a batch with 10 files (not a full batch)
        init_response = self.client.post(
            '/api/applications/bulk-upload/init/',
            content_type='application/json',
            data=json.dumps({'job_listing_id': str(self.job_listing.id)})
        )
        batch_id = init_response.json()['batch_id']

        # Upload 10 files with unique content
        for i in range(10):
            unique_docx = self._create_valid_docx(unique_id=400 + i)
            upload_file = SimpleUploadedFile(
                f"test_resume_{i}.docx",
                unique_docx,
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

            self.client.post(
                '/api/applications/bulk-upload/upload/',
                {
                    'batch_id': batch_id,
                    'file': upload_file
                }
            )

        # Validate
        self.client.post(
            '/api/applications/bulk-upload/validate/',
            content_type='application/json',
            data=json.dumps({'batch_id': batch_id})
        )

        # Commit
        commit_response = self.client.post(
            '/api/applications/bulk-upload/commit/',
            content_type='application/json',
            data=json.dumps({'batch_id': batch_id})
        )

        self.assertEqual(commit_response.status_code, 202)  # Accepted for async processing
        
        # Finalize the batch

        finalize_bulk_upload_batch([], batch_id)

        # Refresh from database
        self.job_listing.refresh_from_db()

        # Verify counters updated
        # batch_count should NOT increment because only 10 files were committed (not 100)
        self.assertEqual(self.job_listing.batch_count, 0)
        # total_resumes should reflect the 10 committed files
        self.assertEqual(self.job_listing.total_resumes, 10)
